"""
One-time setup script for chunking document and creating vector store
Run this script once before using the Streamlit app
"""
import os
from pathlib import Path
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangchainDocument

import config


def load_docx(file_path):
    """Load and extract text from a Word document"""
    print(f"📖 Loading document: {file_path}")
    
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    doc = Document(file_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    text = "\n".join(full_text)
    print(f"✓ Document loaded successfully")
    print(f"✓ Total characters: {len(text):,}")
    
    return text


def chunk_document(text):
    """Split document into chunks"""
    print(f"\n✂️ Splitting document into chunks...")
    print(f"   Chunk size: {config.CHUNK_SIZE}")
    print(f"   Chunk overlap: {config.CHUNK_OVERLAP}")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    print(f"✓ Document split into {len(chunks)} chunks")
    print(f"✓ Average chunk size: {len(text) // len(chunks) if chunks else 0} characters")
    
    return chunks


def create_vector_store(chunks):
    """Create and save FAISS vector store"""
    print(f"\n🔧 Initializing embeddings...")
    
    embeddings = HuggingFaceEmbeddings(
        model_name=config.EMBEDDINGS_MODEL
    )
    print(f"✓ Embeddings model loaded: {config.EMBEDDINGS_MODEL}")
    
    print(f"\n🔨 Creating vector store from {len(chunks)} chunks...")
    
    # Convert chunks to Document objects
    documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
    
    # Create FAISS vector store
    vector_store = FAISS.from_documents(
        documents=documents,
        embedding=embeddings
    )
    
    print(f"✓ Vector store created successfully")
    print(f"✓ Number of vectors: {vector_store.index.ntotal if hasattr(vector_store.index, 'ntotal') else len(documents)}")
    
    # Save vector store
    vector_store_path = config.VECTOR_STORE_PATH / "faiss_index"
    vector_store.save_local(str(vector_store_path))
    print(f"✓ Vector store saved to: {vector_store_path}")
    
    return vector_store


def setup():
    """Run one-time setup"""
    print("\n" + "="*80)
    print("🚀 SETTING UP RAG VECTOR STORE (ONE-TIME SETUP)")
    print("="*80)
    
    # Check if vector store already exists
    vector_store_path = config.VECTOR_STORE_PATH / "faiss_index"
    if (vector_store_path / "index.faiss").exists():
        print("\n⚠️  Vector store already exists!")
        response = input("Do you want to recreate it? (yes/no): ").strip().lower()
        if response != 'yes':
            print("✓ Setup skipped. Using existing vector store.")
            return
    
    try:
        # Step 1: Load document
        document_text = load_docx(config.DOCX_FILE)
        
        # Step 2: Chunk document
        chunks = chunk_document(document_text)
        
        # Step 3: Create and save vector store
        vector_store = create_vector_store(chunks)
        
        print("\n" + "="*80)
        print("✅ SETUP COMPLETED SUCCESSFULLY!")
        print("="*80)
        print(f"✓ Vector store is ready for use in the Streamlit app")
        print(f"✓ Location: {vector_store_path}")
        
    except Exception as e:
        print(f"\n❌ Setup failed with error: {str(e)}")
        raise


if __name__ == "__main__":
    setup()
